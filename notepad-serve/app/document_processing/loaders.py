from langchain_community.document_loaders import (
    TextLoader,
    UnstructuredMarkdownLoader,
    PyPDFLoader,
    DirectoryLoader
)
from langchain_core.documents import Document
from pathlib import Path


class DocumentLoader:
    """文档加载器"""

    LOADERS = {
        ".txt": TextLoader,
        ".md": UnstructuredMarkdownLoader,
        ".pdf": PyPDFLoader,
        ".docx": None,  # 使用自定义加载器
        ".doc": None,   # 使用自定义加载器
    }

    @classmethod
    def _load_docx(cls, file_path: str) -> list:
        """使用 python-docx 加载 Word 文档"""
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            full_text = "\n".join(text_content)
            
            return [Document(
                page_content=full_text,
                metadata={
                    "source": file_path,
                    "file_name": Path(file_path).name,
                    "file_type": Path(file_path).suffix.lower()
                }
            )]
        except Exception as e:
            raise ValueError(f"加载 Word 文档失败: {e}")

    @classmethod
    def load(cls, file_path: str) -> list:
        """加载单个文档"""
        path = Path(file_path)
        suffix = path.suffix.lower()

        if suffix not in cls.LOADERS:
            raise ValueError(f"不支持的文件类型: {suffix}")

        # 使用自定义加载器处理 Word 文档
        if suffix in [".docx", ".doc"]:
            return cls._load_docx(file_path)
        
        loader_class = cls.LOADERS[suffix]
        
        # 对于文本文件，指定 UTF-8 编码
        if suffix == ".txt":
            loader = loader_class(file_path, encoding="utf-8")
        else:
            loader = loader_class(file_path)
        
        documents = loader.load()

        # 添加元数据
        for doc in documents:
            doc.metadata["file_name"] = path.name
            doc.metadata["file_type"] = suffix

        return documents

    @classmethod
    def load_directory(cls, directory: str, glob_pattern: str = "**/*.pdf") -> list:
        """加载目录中的文档"""
        path = Path(directory)
        if not path.exists():
            raise ValueError(f"目录不存在: {directory}")
        
        loader = DirectoryLoader(
            str(directory),
            glob=glob_pattern,
            loader_cls=PyPDFLoader
        )
        documents = loader.load()
        
        # 添加元数据
        for doc in documents:
            file_path = Path(doc.metadata.get("source", ""))
            doc.metadata["file_name"] = file_path.name
            doc.metadata["file_type"] = file_path.suffix.lower()
        
        return documents

    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """检查是否支持该文件类型"""
        suffix = Path(file_path).suffix.lower()
        return suffix in cls.LOADERS
